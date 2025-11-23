# -*- coding: utf-8 -*-
"""
Модуль RAG-системы с контекстной памятью
=========================================

Этот модуль реализует улучшенную RAG (Retrieval-Augmented Generation) систему
с поддержкой контекстной памяти диалога. Основные возможности:

- Загрузка и обработка PDF/DOCX документов
- Создание и управление векторной базой знаний
- Семантический поиск по документам
- Генерация ответов с учетом истории диалога
- Контекстное переформулирование запросов

Архитектура системы:
1. DocumentProcessor - обработка документов
2. VectorStoreManager - управление векторной БД  
3. MemoryManager - управление контекстом диалога
4. QueryProcessor - обработка запросов и генерация ответов
"""

import os
import re
from loguru import logger

# Импорт библиотек LangChain
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_ollama import ChatOllama
from langchain_core.messages import HumanMessage
from langchain.memory import ConversationBufferWindowMemory

# =============================================================================
# КОНФИГУРАЦИЯ И НАСТРОЙКИ
# =============================================================================

# Настройки векторной базы данных
VECTOR_DB_PATH = 'db/db_01'
PROCESSED_FILES_LOG = 'db/db_01/processed_files.txt'

# Настройки модели эмбеддингов
EMBEDDING_MODEL = 'intfloat/multilingual-e5-large'
EMBEDDING_DEVICE = 'cpu'

# Настройки языковой модели
LLM_MODEL = "qwen2.5-coder:7b"
LLM_TEMPERATURE = 0

# Настройки обработки документов
CHUNK_SIZE = 1024
CHUNK_OVERLAP = 0
DEFAULT_RELEVANT_CHUNKS = 3

# Настройки памяти диалога
MEMORY_WINDOW_SIZE = 6  # Сохраняем последние 3 пары вопрос-ответ

# =============================================================================
# НАСТРОЙКА ЛОГГИРОВАНИЯ
# =============================================================================

logger.add(
    "log/02_Simple_RAG_PDF.log", 
    format="{time:YYYY-MM-DD HH:mm:ss} {level} {message}",
    level="DEBUG", 
    rotation="100 KB", 
    compression="zip",
    retention="10 days"
)

# =============================================================================
# МЕНЕДЖЕР ПАМЯТИ СЕССИИ
# =============================================================================

# Глобальная переменная для хранения памяти сессии
_session_memory = None

def get_session_memory():
    """
    Инициализация или получение текущей сессионной памяти.
    
    Returns:
        ConversationBufferWindowMemory: Объект памяти с историей диалога
    """
    global _session_memory
    if _session_memory is None:
        _session_memory = ConversationBufferWindowMemory(
            memory_key="chat_history",
            return_messages=True,
            output_key="answer",
            k=MEMORY_WINDOW_SIZE
        )
        logger.debug("Инициализирована новая сессионная память")
    return _session_memory

def clear_session_memory():
    """
    Очистка сессионной памяти для начала нового диалога.
    """
    global _session_memory
    _session_memory = None
    logger.debug("Сессионная память очищена")

def get_memory_context():
    """
    Получение форматированной истории диалога из памяти.
    
    Returns:
        str: Отформатированная история диалога или пустая строка
    """
    memory = get_session_memory()
    history_dict = memory.load_memory_variables({})
    
    if "chat_history" not in history_dict or not history_dict["chat_history"]:
        return ""
    
    # Форматируем историю диалога для читаемости
    chat_history_parts = []
    for i, msg in enumerate(history_dict["chat_history"]):
        role = "Пользователь" if i % 2 == 0 else "Ассистент"
        content = msg.content if hasattr(msg, 'content') else str(msg)
        chat_history_parts.append(f"{role}: {content}")
    
    return "\n".join(chat_history_parts)

# =============================================================================
# МЕНЕДЖЕР ВЕКТОРНОЙ БАЗЫ ДАННЫХ
# =============================================================================

def initialize_embeddings():
    """
    Инициализация модели для создания векторных представлений текста.
    
    Returns:
        HuggingFaceEmbeddings: Модель для создания эмбеддингов
    """
    return HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={'device': EMBEDDING_DEVICE}
    )

def get_current_document_files():
    """
    Сканирование директории с документами для получения актуального списка файлов.
    
    Returns:
        list: Список путей к PDF и DOCX файлам
    """
    current_files = []
    for root, dirs, files in os.walk('pdf'):
        for file in files:
            if file.endswith((".pdf", ".docx")):
                current_files.append(os.path.join(root, file))
    return current_files

def load_processed_files():
    """
    Загрузка списка ранее обработанных файлов из метаданных.
    
    Returns:
        set: Множество путей к обработанным файлам
    """
    if not os.path.exists(PROCESSED_FILES_LOG):
        return set()
    
    with open(PROCESSED_FILES_LOG, 'r', encoding='utf-8') as f:
        return set(line.strip() for line in f.readlines())

def save_processed_files(file_paths):
    """
    Сохранение списка обработанных файлов в метаданные.
    
    Args:
        file_paths (list): Список путей к обработанным файлам
    """
    os.makedirs(VECTOR_DB_PATH, exist_ok=True)
    with open(PROCESSED_FILES_LOG, 'w', encoding='utf-8') as f:
        for file_path in file_paths:
            f.write(file_path + '\n')
    logger.debug(f"Список обработанных файлов сохранен в {PROCESSED_FILES_LOG}")

def should_rebuild_database():
    """
    Проверка необходимости перестроения векторной базы данных.
    
    Returns:
        tuple: (need_rebuild, current_files, processed_files)
    """
    index_path = VECTOR_DB_PATH + "/index.faiss"
    
    if not os.path.exists(index_path):
        logger.debug('Векторная база не существует, требуется создание')
        return True, [], set()
    
    logger.debug('Векторная база существует, проверка актуальности')
    current_files = get_current_document_files()
    processed_files = load_processed_files()
    
    if set(current_files) == processed_files:
        logger.debug('Изменений в документах не обнаружено')
        return False, current_files, processed_files
    else:
        logger.debug('Обнаружены изменения в документах, требуется перестроение')
        return True, current_files, processed_files

def load_documents_from_directory(directory='pdf'):
    """
    Загрузка и обработка документов из указанной директории.
    
    Args:
        directory (str): Путь к директории с документами
        
    Returns:
        tuple: (documents, file_paths) - список документов и путей к файлам
    """
    documents = []
    file_paths = []
    
    logger.debug(f'Сканирование директории: {directory}')
    
    for root, dirs, files in os.walk(directory):
        for file in files:
            file_path = os.path.join(root, file)
            
            if file.endswith(".pdf"):
                logger.debug(f'Обработка PDF: {file_path}')
                try:
                    loader = PyPDFLoader(file_path)
                    documents.extend(loader.load())
                    file_paths.append(file_path)
                except Exception as e:
                    logger.error(f"Ошибка загрузки PDF {file_path}: {e}")
                    
            elif file.endswith(".docx"):
                logger.debug(f'Обработка DOCX: {file_path}')
                try:
                    loader = Docx2txtLoader(file_path)
                    docs = loader.load()
                    documents.extend(docs)
                    file_paths.append(file_path)
                    logger.debug(f"Успешно загружен DOCX: {file_path}")
                except Exception as e:
                    logger.error(f"Ошибка загрузки DOCX {file_path}: {e}")
                    # Резервный метод загрузки DOCX
                    documents.extend(load_docx_fallback(file_path))
                    file_paths.append(file_path)

    return documents, file_paths

def load_docx_fallback(file_path):
    """
    Резервный метод загрузки DOCX файлов через библиотеку python-docx.
    
    Args:
        file_path (str): Путь к DOCX файлу
        
    Returns:
        list: Список документов LangChain
    """
    try:
        import docx
        from langchain_core.documents import Document
        
        doc = docx.Document(file_path)
        full_text = [paragraph.text for paragraph in doc.paragraphs]
        doc_content = '\n'.join(full_text)
        
        logger.debug(f"DOCX загружен через python-docx: {file_path}")
        return [Document(page_content=doc_content, metadata={"source": file_path})]
    except Exception as e:
        logger.error(f"Критическая ошибка загрузки {file_path}: {e}")
        return []

def split_documents(documents):
    """
    Разделение документов на чанки для векторного поиска.
    
    Args:
        documents (list): Список документов LangChain
        
    Returns:
        list: Список разделенных чанков
    """
    logger.debug('Разделение документов на чанки')
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP
    )
    
    chunks = text_splitter.split_documents(documents)
    
    logger.debug(f'Тип чанков: {type(chunks)}')
    logger.debug(f'Количество чанков: {len(chunks)}')
    
    # Логирование для отладки
    if chunks:
        sample_idx = min(100, len(chunks) - 1)
        logger.debug(f"Пример чанка {sample_idx}: {chunks[sample_idx].metadata}")
    
    return chunks

def get_index_db():
    """
    Основная функция для получения или создания векторной базы знаний.
    
    Returns:
        FAISS: Объект векторной базы данных
    """
    logger.debug('Инициализация векторной базы данных...')
    
    embeddings = initialize_embeddings()
    need_rebuild, current_files, processed_files = should_rebuild_database()
    
    if not need_rebuild:
        logger.debug('Загрузка существующей векторной базы')
        return FAISS.load_local(VECTOR_DB_PATH, embeddings, allow_dangerous_deserialization=True)

    # Процесс перестроения базы данных
    logger.debug('Создание/обновление векторной базы знаний')
    
    documents, file_paths = load_documents_from_directory()
    
    if not documents:
        logger.warning('Не найдено документов для обработки!')
        from langchain_core.documents import Document
        documents = [Document(page_content="Нет данных", metadata={})]
    
    # Разделение на чанки и создание векторной базы
    chunks = split_documents(documents)
    db = FAISS.from_documents(chunks, embeddings)
    
    # Сохранение базы и метаданных
    logger.debug('Сохранение векторной базы знаний')
    db.save_local(VECTOR_DB_PATH)
    save_processed_files(file_paths)
    
    logger.debug('Векторная база успешно создана/обновлена')
    return db

# =============================================================================
# ОБРАБОТЧИК ЗАПРОСОВ И КОНТЕКСТА
# =============================================================================

def get_contextual_query(topic, memory):
    """
    Переформулирование запроса с учетом контекста предыдущего диалога.
    
    Args:
        topic (str): Оригинальный запрос пользователя
        memory: Объект памяти с историей диалога
        
    Returns:
        str: Улучшенный запрос с контекстом
    """
    if not memory:
        return topic
    
    # Получаем историю диалога
    history_dict = memory.load_memory_variables({})
    if "chat_history" not in history_dict or len(history_dict["chat_history"]) < 2:
        return topic
    
    # Извлекаем последнюю пару вопрос-ответ
    last_interaction = history_dict["chat_history"][-2:]
    last_question = last_interaction[0].content if hasattr(last_interaction[0], 'content') else str(last_interaction[0])
    last_answer = last_interaction[1].content if hasattr(last_interaction[1], 'content') else str(last_interaction[1])
    
    # Создаем расширенный запрос с контекстом
    contextual_query = (
        f"{topic} [контекст: предыдущий вопрос: '{last_question}', "
        f"предыдущий ответ: '{last_answer}']"
    )
    
    logger.debug(f"Оригинальный запрос: {topic}")
    logger.debug(f"Контекстный запрос: {contextual_query}")
    
    return contextual_query

def get_message_content(topic, db, number_relevant_chunks, memory=None):
    """
    Извлечение релевантных фрагментов из векторной базы знаний.
    
    Args:
        topic (str): Запрос пользователя
        db: Векторная база данных FAISS
        number_relevant_chunks (int): Количество извлекаемых фрагментов
        memory: Объект памяти для контекстного поиска
        
    Returns:
        str: Отформатированное содержимое релевантных фрагментов
    """
    logger.debug('Поиск релевантных фрагментов...')
    
    # Используем контекстный запрос если доступна память
    search_query = get_contextual_query(topic, memory) if memory else topic
    
    # Выполняем семантический поиск
    docs = db.similarity_search(search_query, k=number_relevant_chunks)
    
    # Форматируем результаты
    formatted_chunks = []
    for i, doc in enumerate(docs):
        chunk_info = (
            f'\n#### Фрагмент {i+1} ####\n'
            f'Метаданные: {doc.metadata}\n'
            f'Содержимое: {doc.page_content}\n'
        )
        formatted_chunks.append(chunk_info)
    
    message_content = re.sub(r'\n{2}', ' ', '\n '.join(formatted_chunks))
    logger.debug(f"Найдено релевантных фрагментов: {len(docs)}")
    
    return message_content

# =============================================================================
# ГЕНЕРАТОР ОТВЕТОВ
# =============================================================================

def initialize_llm():
    """
    Инициализация языковой модели для генерации ответов.
    
    Returns:
        ChatOllama: Объект языковой модели
    """
    return ChatOllama(
        model=LLM_MODEL,
        temperature=LLM_TEMPERATURE
    )

def create_prompt_template(has_history=False):
    """
    Создание шаблона промпта в зависимости от наличия истории диалога.
    
    Args:
        has_history (bool): Флаг наличия истории диалога
        
    Returns:
        str: Шаблон промпта
    """
    if has_history:
        return """Ты являешься помощником для ответов на вопросы. Внимательно проанализируй всю историю диалога и текущий вопрос.

ИСТОРИЯ ДИАЛОГА:
{history}

ТЕКУЩИЙ ВОПРОС ПОЛЬЗОВАТЕЛЯ: 
{question}

КОНТЕКСТ ДЛЯ ОТВЕТА (из документов):
{context}

ИНСТРУКЦИИ:
1. Внимательно проанализируй историю диалога и текущий вопрос
2. Если текущий вопрос является уточнением или продолжением предыдущих вопросов - учти это в ответе
3. Используй информацию из контекста для ответа
4. Будь последовательным и логичным в своих ответах
5. Если вопрос явно связан с предыдущей темой, сделай плавный переход
6. Ответ должен быть лаконичным, но полным

Ответ:"""
    else:
        return """Ты являешься помощником для выполнения заданий по ответам на вопросы. 
Используй следующий контекст для ответа на вопрос:
{context} 

Вопрос: {question}

Дай точный и лаконичный ответ на вопрос, используя только предоставленный контекст.
Ответ:"""

def get_model_response(topic, message_content, use_memory=True):
    """
    Генерация ответа модели на основе контекста и истории диалога.
    
    Args:
        topic (str): Запрос пользователя
        message_content (str): Релевантный контекст из документов
        use_memory (bool): Флаг использования памяти диалога
        
    Returns:
        str: Сгенерированный ответ модели
    """
    logger.debug('Генерация ответа модели...')

    # Инициализация моделей
    llm = initialize_llm()
    memory = get_session_memory() if use_memory else None
    
    # Получение истории диалога
    chat_history = get_memory_context() if use_memory and memory else ""
    
    # Выбор шаблона промпта
    prompt_template = create_prompt_template(bool(chat_history))
    
    # Форматирование промпта
    if chat_history:
        formatted_prompt = prompt_template.format(
            context=message_content, 
            history=chat_history, 
            question=topic
        )
    else:
        formatted_prompt = prompt_template.format(
            context=message_content, 
            question=topic
        )
    
    # Генерация ответа
    generation = llm.invoke([HumanMessage(content=formatted_prompt)])
    model_response = generation.content
    
    # Сохранение в память диалога
    if use_memory and memory:
        memory.save_context({"input": topic}, {"answer": model_response})
        logger.debug("Диалог сохранен в память сессии")
    
    logger.debug(f"Сгенерирован ответ длиной {len(model_response)} символов")
    return model_response

# =============================================================================
# ИНТЕРФЕЙСНЫЕ ФУНКЦИИ
# =============================================================================

def chat_with_memory(topic, db, number_relevant_chunks=3, use_memory=True):
    """
    Упрощенный интерфейс для чата с поддержкой памяти.
    
    Args:
        topic (str): Запрос пользователя
        db: Векторная база данных
        number_relevant_chunks (int): Количество релевантных фрагментов
        use_memory (bool): Флаг использования памяти
        
    Returns:
        str: Ответ системы
    """
    logger.debug(f"Обработка запроса: {topic}")
    
    memory = get_session_memory() if use_memory else None
    message_content = get_message_content(topic, db, number_relevant_chunks, memory)
    response = get_model_response(topic, message_content, use_memory)
    
    return response

def interactive_chat_loop():
    """
    Интерактивный режим чата для тестирования системы.
    """
    print("=== Улучшенная RAG-система с контекстной памятью ===")
    print("Система загружается...")
    
    # Инициализация базы данных
    db = get_index_db()
    
    print("Система готова к работе!")
    print("Доступные команды:")
    print("  'clear' - очистить память диалога")
    print("  'quit', 'exit', 'выход' - завершить работу")
    print("-" * 50)
    
    while True:
        try:
            user_input = input("\nВаш вопрос: ").strip()
            
            if not user_input:
                continue
                
            if user_input.lower() in ['quit', 'exit', 'выход']:
                print("Завершение работы...")
                break
                
            if user_input.lower() == 'clear':
                clear_session_memory()
                print("Память диалога очищена")
                continue
            
            # Обработка запроса пользователя
            response = chat_with_memory(user_input, db)
            print(f"Ответ: {response}")
            
        except KeyboardInterrupt:
            print("\n\nЗавершение работы по запросу пользователя...")
            break
        except Exception as e:
            logger.error(f"Ошибка в основном цикле: {e}")
            print("⚠️ Произошла ошибка. Попробуйте еще раз.")

# =============================================================================
# ТОЧКА ВХОДА
# =============================================================================

if __name__ == "__main__":
    
    # Запуск интерактивного режима
    interactive_chat_loop()